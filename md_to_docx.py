# -*- coding: utf-8 -*-
"""
把商业计划书_鲜智巡检.md 转换为 Word 文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def set_font(run, size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    sizes = {1: 22, 2: 18, 3: 15, 4: 13}
    set_font(run, size=sizes.get(level, 12), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, size=11):
    p = doc.add_paragraph()
    # 处理行内加粗 **text**
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_font(run, size=size, bold=(i % 2 == 1))
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # 灰色背景效果用缩进模拟
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_from_md(doc, lines):
    """解析markdown表格并插入word表格"""
    rows = []
    for line in lines:
        if re.match(r'\|[-| :]+\|', line):
            continue  # 跳过分隔行
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= col_count:
                break
            cell = table.cell(i, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            set_font(run, size=10, bold=(i == 0))
    doc.add_paragraph()

def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 页面设置：A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    i = 0
    in_code = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 代码块
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, '\n'.join(code_lines))
                in_code = False
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # 表格
        if line.strip().startswith('|'):
            table_lines.append(line.strip())
            i += 1
            # 收集连续的表格行
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_table_from_md(doc, table_lines)
            table_lines = []
            continue

        # 标题
        if line.startswith('#### '):
            add_heading(doc, line[5:], 4)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
        elif line.startswith('# '):
            add_heading(doc, line[2:], 1)

        # 分隔线
        elif line.strip() == '---':
            doc.add_paragraph('─' * 40)

        # 引用块
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_font(run, size=10, color=(100, 100, 100))
            p.paragraph_format.left_indent = Cm(1)

        # 无序列表
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line.strip()[2:]
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                set_font(run, size=11, bold=(j % 2 == 1))

        # 空行
        elif line.strip() == '':
            pass

        # 普通段落
        else:
            # 去掉图片占位符的**加粗**，保留文字
            text = line.strip()
            if text:
                add_paragraph(doc, text)

        i += 1

    doc.save(docx_path)
    print(f"✅ 转换完成：{docx_path}")

if __name__ == '__main__':
    convert('商业计划书_鲜智巡检.md', '商业计划书_鲜智巡检.docx')
